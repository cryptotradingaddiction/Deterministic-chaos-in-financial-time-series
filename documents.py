#!/usr/bin/env python3
"""
Build thesis-style summary tables into a Word document for the DCh pipeline.

**When it runs**

``hypothesis.bat`` invokes this script after the three invariant batches finish.
It can also be run manually: ``py -3 documents.py``.

**Output**

``<paths.results_dir>/results.docx`` (landscape A4, Arial 8pt tables).

**Data sources**

- Log-return CSVs (liquidity-cut when present) — sample dates and *N*.
- ``tau_w/_tau_w_summary.txt``, ``2dc/_2dc_summary.txt``, ``mutual/`` (τ via hypothesis agg).
- ``theiler_w/_theiler_summary.txt`` or ``W_D2_<sym>`` in ``_per_coin_settings.bat``.
- ``_hypothesis_aggregate_summary.txt`` under ``correlation_dimension_*``,
  ``lambda_max_*``, and ``rqa_*`` (original invariants + bootstrap TS columns).

**Tables in the document**

1. Wide invariant matrix (one row per quantity, columns = coins).
2. Surrogate / bootstrap decision table (TS rule, |TS| > 3).
3. Input series statistics (μ, σ for orig / surr / normal / t3.5).
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import os
import re
import pandas as pd

from config_loader import (
    get_data_dir,
    get_results_dir,
    load_config,
    prefer_liquidity_cut,
    hypothesis_correlation_dimension_dir,
    hypothesis_lambda_max_dir,
    hypothesis_rqa_dir,
    dch_test_mode_from_env,
    default_per_coin_settings_bat_path,
    parse_per_coin_settings_bat,
    parse_theiler_w_map,
    theiler_summary_path,
)

try:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
except ModuleNotFoundError as exc:  # pragma: no cover
    raise SystemExit(
        "Missing dependency: python-docx. Install with: py -3 -m pip install python-docx"
    ) from exc


def _results_root() -> Path:
    """``paths.results_dir`` from config (hypothesis aggregates live here)."""
    return Path(get_results_dir(load_config()))


def _out_docx_path() -> Path:
    """Fixed output filename at the results root."""
    return _results_root() / "results.docx"


# Coins and Word table column order (Czech headers for thesis layout).
SYMS = ["BTCUSD", "ETHUSD", "XRPUSD", "LTCUSD", "LINKUSD", "DOGEUSD", "ADAUSD"]
HEADERS = ["Aktivum / USD", "BTC", "ETH", "XRP", "LTC", "LINK", "DOGE", "ADA"]


@dataclass
class TableData:
    """One string cell per symbol for the main invariant summary table."""

    sample_start: dict[str, str]
    sample_end: dict[str, str]
    n_vals: dict[str, str]
    tau: dict[str, str]
    tau_w: dict[str, str]
    theiler_w: dict[str, str]
    m_vals: dict[str, str]
    dc: dict[str, str]
    takens: dict[str, str]
    ellner: dict[str, str]
    lle: dict[str, str]
    rr: dict[str, str]
    det: dict[str, str]
    lam: dict[str, str]
    maxline: dict[str, str]
    entr: dict[str, str]
    tt: dict[str, str]
    trend: dict[str, str]


@dataclass
class SeriesStatsRow:
    """One row in the input log-return statistics table (means/SDs of the series, not invariants)."""

    symbol: str
    orig_mean: str
    orig_sd: str
    surr_mean: str
    surr_sd: str
    normal_mean: str
    normal_sd: str
    t_mean: str
    t_sd: str


def _parse_agg_col(path: Path, col_index: int, decimals: int = 4, as_int: bool = False) -> dict[str, str]:
    """
    Parse whitespace-separated aggregate files by **column index** (0-based after symbol).

    Used for legacy summaries without a ``Symbol ...`` header row.
    """
    txt = path.read_text(encoding="utf-8", errors="ignore")
    out: dict[str, str] = {}
    for line in txt.splitlines():
        m = re.match(r"^(ADAUSD|BTCUSD|DOGEUSD|ETHUSD|LINKUSD|LTCUSD|XRPUSD)\s+", line.strip())
        if not m:
            continue
        parts = line.split()
        sym = parts[0]
        val = float(parts[col_index])
        if as_int:
            out[sym] = str(int(round(val)))
        else:
            out[sym] = f"{val:.{decimals}f}"
    return out


def _parse_agg_named_col(path: Path, col_name: str, decimals: int = 4, as_int: bool = False) -> dict[str, str]:
    """
    Parse ``_hypothesis_aggregate_summary.txt``-style files by **column name**.

    Expects a header line starting with ``Symbol ``; data rows start with a known sym.
    """
    txt = path.read_text(encoding="utf-8", errors="ignore")
    header: list[str] | None = None
    out: dict[str, str] = {}
    for line in txt.splitlines():
        stripped = line.strip()
        if stripped.startswith("Symbol "):
            header = stripped.split()
            continue
        m = re.match(r"^(ADAUSD|BTCUSD|DOGEUSD|ETHUSD|LINKUSD|LTCUSD|XRPUSD)\s+", stripped)
        if not m or header is None:
            continue
        parts = stripped.split()
        if col_name not in header or len(parts) <= header.index(col_name):
            continue
        sym = parts[0]
        val = float(parts[header.index(col_name)])
        if as_int:
            out[sym] = str(int(round(val)))
        else:
            out[sym] = f"{val:.{decimals}f}"
    return out


def _format_float_token(token: str, decimals: int = 4) -> str:
    """Format one aggregate field for Word; empty string on parse failure."""
    try:
        val = float(token)
    except (TypeError, ValueError):
        return ""
    if pd.isna(val):
        return "nan"
    return f"{val:.{decimals}f}"


def _parse_aggregate_rows(path: Path) -> tuple[list[str], dict[str, dict[str, str]]]:
    """Return (header_names, {symbol: {col_name: token}}) from a hypothesis aggregate file."""
    txt = path.read_text(encoding="utf-8", errors="ignore")
    header: list[str] = []
    rows: dict[str, dict[str, str]] = {}
    for line in txt.splitlines():
        stripped = line.strip()
        if stripped.startswith("Symbol "):
            header = stripped.split()
            continue
        m = re.match(r"^(ADAUSD|BTCUSD|DOGEUSD|ETHUSD|LINKUSD|LTCUSD|XRPUSD)\s+", stripped)
        if not m or not header:
            continue
        parts = stripped.split()
        sym = parts[0]
        rows[sym] = {name: parts[i] for i, name in enumerate(header) if i < len(parts)}
    return header, rows


@dataclass
class SurrogateRow:
    """One row in the surrogate/bootstrap results table (per coin × invariant)."""

    symbol: str
    invariant: str
    orig: str
    boot_mean: str
    boot_sd: str
    resh: str
    ts: str
    abs_ts: str
    decision: str


def _bootstrap_decision(abs_ts: str, threshold: float = 3.0) -> str:
    """Map |TS| to Czech decision text; ``not bootstrap-tested`` when TS is missing."""
    try:
        v = float(abs_ts)
    except (TypeError, ValueError):
        return "not bootstrap-tested"
    if pd.isna(v):
        return "not bootstrap-tested"
    return "reject H0" if v > threshold else "fail to reject H0"


def _collect_surrogate_rows(config=None) -> list[SurrogateRow]:
    """
    Flatten hypothesis aggregates into surrogate table rows for all coins × metrics.

    Reads ``_hypothesis_aggregate_summary.txt`` from dimension, LLE, and RQA result
    trees (FULL or test_<N> via :func:`config_loader.hypothesis_*_dir`).
    """
    cfg = config or load_config()
    dim_agg = hypothesis_correlation_dimension_dir(cfg) / "_hypothesis_aggregate_summary.txt"
    lle_agg = hypothesis_lambda_max_dir(cfg) / "_hypothesis_aggregate_summary.txt"
    rqa_agg = hypothesis_rqa_dir(cfg) / "_hypothesis_aggregate_summary.txt"
    sources = {
        "TAKENS": dim_agg,
        "ELLNER": dim_agg,
        "LLE": lle_agg,
        "RR": rqa_agg,
        "DET": rqa_agg,
        "LAM": rqa_agg,
        "MAXLINE": rqa_agg,
        "ENTR": rqa_agg,
        "TT": rqa_agg,
        "TREND": rqa_agg,
    }
    metric_order = ["TAKENS", "ELLNER", "LLE", "RR", "DET", "LAM", "MAXLINE", "ENTR", "TT", "TREND"]
    label = {"TAKENS": "D_T (Takens)", "ELLNER": "D_E (Ellner)"}
    parsed = {}
    for path in set(sources.values()):
        if path.exists():
            _header, rows = _parse_aggregate_rows(path)
            parsed[path] = rows
    out: list[SurrogateRow] = []
    for sym in SYMS:
        for metric in metric_order:
            path = sources[metric]
            row = parsed.get(path, {}).get(sym, {})
            abs_ts = _format_float_token(row.get(f"absTS_{metric}", "nan"))
            out.append(
                SurrogateRow(
                    symbol=sym.replace("USD", ""),
                    invariant=label.get(metric, metric),
                    orig=_format_float_token(row.get(f"{metric}_orig", "nan")),
                    boot_mean=_format_float_token(row.get(f"{metric}_boot", "nan")),
                    boot_sd=_format_float_token(row.get(f"{metric}_boot_sd", "nan")),
                    resh=_format_float_token(row.get(f"{metric}_resh", "nan")),
                    ts=_format_float_token(row.get(f"TS_{metric}", "nan")),
                    abs_ts=abs_ts,
                    decision=_bootstrap_decision(abs_ts),
                )
            )
    return out


def _parse_series_stats_from_summary(path: Path) -> dict[str, tuple[str, str]]:
    """
    Extract μ and σ for orig / surr / normal / t3.5 from one ``*_surrogate_summary.txt``.

    Looks for the ``Series statistics`` block written by ``hypothesis.py``.
    """
    txt = path.read_text(encoding="utf-8", errors="ignore")
    out: dict[str, tuple[str, str]] = {}
    in_block = False
    for line in txt.splitlines():
        if line.strip().startswith("Series statistics"):
            in_block = True
            continue
        if in_block and line.strip().startswith("Invariant"):
            break
        m = re.match(r"^\s*(orig|surr|normal|t3\.5)\s+(-?\d+\.\d+)\s+(-?\d+\.\d+)\s*$", line)
        if m:
            out[m.group(1)] = (m.group(2), m.group(3))
    return out


def _collect_series_stats_rows(config=None) -> list[SeriesStatsRow]:
    # The input-series statistics are properties of the generated time series,
    # not of individual invariants. Use the D2/Takens summaries as one
    # representative source to avoid repeating identical rows per invariant.
    cfg = config or load_config()
    root = hypothesis_correlation_dimension_dir(cfg)
    per_symbol: dict[str, dict[str, tuple[str, str]]] = {}
    if root.exists():
        for fp in root.rglob("*_surrogate_summary.txt"):
            sym = fp.name.replace("_surrogate_summary.txt", "")
            if sym in SYMS:
                per_symbol[sym] = _parse_series_stats_from_summary(fp)
    out: list[SeriesStatsRow] = []
    for sym in SYMS:
        stats = per_symbol.get(sym, {})
        orig = stats.get("orig", ("nan", "nan"))
        surr = stats.get("surr", ("nan", "nan"))
        normal = stats.get("normal", ("nan", "nan"))
        t_ref = stats.get("t3.5", ("nan", "nan"))
        out.append(
            SeriesStatsRow(
                symbol=sym.replace("USD", ""),
                orig_mean=orig[0],
                orig_sd=orig[1],
                surr_mean=surr[0],
                surr_sd=surr[1],
                normal_mean=normal[0],
                normal_sd=normal[1],
                t_mean=t_ref[0],
                t_sd=t_ref[1],
            )
        )
    return out


def _parse_tau_map(config=None) -> dict[str, str]:
    """
    Embedding delay τ per symbol for the Word table.

    Prefer the ``tau`` column in ``_hypothesis_aggregate_summary.txt`` (same values
    passed to ``hypothesis.py --delay``). Fall back to ``TAU_D2_<sym>`` in
    ``_per_coin_settings.bat`` when the aggregate is missing or empty.
    """
    cfg = config or load_config()
    path = hypothesis_correlation_dimension_dir(cfg) / "_hypothesis_aggregate_summary.txt"
    out: dict[str, str] = {}
    if path.is_file():
        out = _parse_agg_named_col(path, "tau", as_int=True)
    if len(out) < len(SYMS):
        bat = parse_per_coin_settings_bat(default_per_coin_settings_bat_path())
        sk = {k.upper(): v for k, v in bat.items()}
        for sym in SYMS:
            if sym in out and out[sym]:
                continue
            raw = sk.get(f"TAU_D2_{sym}")
            if raw is not None:
                try:
                    out[sym] = str(int(float(raw)))
                except (TypeError, ValueError):
                    pass
    return out


def _parse_tauw_map() -> dict[str, str]:
    """Decorrelation time τ_w from ``tau_w/_tau_w_summary.txt`` (fixed-width table)."""
    df = pd.read_fwf(_results_root() / "tau_w" / "_tau_w_summary.txt")
    out: dict[str, str] = {}
    for _, r in df.iterrows():
        sym = str(r["Symbol"]).strip()
        if sym in SYMS:
            out[sym] = f"{float(r['final_tau_w']):.4f}"
    return out


def _theiler_w_for_doc(config=None) -> dict[str, str]:
    """Theiler W for Word tables — equals embedding delay τ (``W_D2_<sym> := TAU_D2_<sym>``)."""
    return _parse_tau_map(config)


def _parse_dc_map() -> dict[str, str]:
    """Capacity dimension d_c from ``2dc/_2dc_summary.txt``."""
    df = pd.read_fwf(_results_root() / "2dc" / "_2dc_summary.txt")
    out: dict[str, str] = {}
    for _, r in df.iterrows():
        sid = str(r["series_id"]).strip()
        sym = sid.split("_")[0]
        if sym in SYMS:
            out[sym] = f"{float(r['d_c']):.4f}"
    return out


def _parse_start_end_n() -> tuple[dict[str, str], dict[str, str], dict[str, str]]:
    """
    Sample window metadata from liquidity-cut log-return CSVs.

    Returns ``(start_dates, end_dates, n_rows)`` keyed by symbol; dates formatted
    ``dd.mm.yyyy, HH:MM`` for the thesis table.
    """
    starts: dict[str, str] = {}
    ends: dict[str, str] = {}
    nvals: dict[str, str] = {}
    cfg = load_config()
    data_dir = get_data_dir(cfg)
    for sym in SYMS:
        raw = os.path.join(data_dir, f"{sym}_BITSTAMP_1h_complete_logreturns.csv")
        path = prefer_liquidity_cut(raw)
        df = pd.read_csv(path)
        d = pd.to_datetime(df["datetime_str"], format="%Y-%m-%d %H:%M:%S", errors="coerce")
        d = d.dropna()
        starts[sym] = d.iloc[0].strftime("%d.%m.%Y, %H:%M")
        ends[sym] = d.iloc[-1].strftime("%d.%m.%Y, %H:%M")
        nvals[sym] = str(len(df))
    return starts, ends, nvals


def _parse_takens(config=None) -> dict[str, str]:
    """Original-series TAKENS plateau mean (D_T) from dimension aggregate."""
    cfg = config or load_config()
    return _parse_agg_named_col(
        hypothesis_correlation_dimension_dir(cfg) / "_hypothesis_aggregate_summary.txt",
        "TAKENS_orig",
        decimals=4,
    )


def _parse_ellner(config=None) -> dict[str, str]:
    """Original-series ELLNER extension (D_E) from dimension aggregate."""
    cfg = config or load_config()
    return _parse_agg_named_col(
        hypothesis_correlation_dimension_dir(cfg) / "_hypothesis_aggregate_summary.txt",
        "ELLNER_orig",
        decimals=4,
    )


def _parse_lle(config=None) -> dict[str, str]:
    """Original-series LLE (λ_max proxy) from LLE aggregate."""
    cfg = config or load_config()
    return _parse_agg_named_col(
        hypothesis_lambda_max_dir(cfg) / "_hypothesis_aggregate_summary.txt",
        "LLE_orig",
        decimals=4,
    )


def _parse_rqa_row(metric_col: str, decimals: int = 4, config=None) -> dict[str, str]:
    """One RQA scalar column from the RQA hypothesis aggregate (orig values)."""
    cfg = config or load_config()
    col_name = {
        "RR_orig_mean": "RR_orig",
        "DET_orig_mean": "DET_orig",
        "LAM_orig_mean": "LAM_orig",
        "MAXLINE_orig_mean": "MAXLINE_orig",
        "ENTR_orig_mean": "ENTR_orig",
        "TT_orig_mean": "TT_orig",
        "TREND_orig_mean": "TREND_orig",
    }[metric_col]
    return _parse_agg_named_col(
        hypothesis_rqa_dir(cfg) / "_hypothesis_aggregate_summary.txt",
        col_name,
        decimals=decimals,
        as_int=(metric_col == "MAXLINE_orig_mean"),
    )


def collect(config=None) -> TableData:
    """Gather all per-symbol fields for the main invariant Word table."""
    cfg = config or load_config()
    starts, ends, nvals = _parse_start_end_n()
    return TableData(
        sample_start=starts,
        sample_end=ends,
        n_vals=nvals,
        tau=_parse_tau_map(cfg),
        tau_w=_parse_tauw_map(),
        theiler_w=_theiler_w_for_doc(cfg),
        m_vals={s: "3" for s in SYMS},  # fixed embedding dim in active pipeline
        dc=_parse_dc_map(),
        takens=_parse_takens(cfg),
        ellner=_parse_ellner(cfg),
        lle=_parse_lle(cfg),
        rr=_parse_rqa_row("RR_orig_mean", config=cfg),
        det=_parse_rqa_row("DET_orig_mean", config=cfg),
        lam=_parse_rqa_row("LAM_orig_mean", config=cfg),
        maxline=_parse_rqa_row("MAXLINE_orig_mean", config=cfg),
        entr=_parse_rqa_row("ENTR_orig_mean", config=cfg),
        tt=_parse_rqa_row("TT_orig_mean", config=cfg),
        trend=_parse_rqa_row("TREND_orig_mean", config=cfg),
    )


def row_values(label: str, mapping: dict[str, str]) -> list[str]:
    """Build one table row: Czech label + values in SYMS column order."""
    return [label] + [mapping.get(sym, "") for sym in SYMS]


# ---------------------------------------------------------------------------
# python-docx formatting helpers
# ---------------------------------------------------------------------------


def _set_cell_text(cell, text: str, font_size: float = 8.0, bold: bool = False) -> None:
    """Centre text in a table cell with uniform font size."""
    cell.text = str(text)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.size = Pt(font_size)
            run.bold = bold


def _format_table(table, widths: list[float], font_size: float = 8.0) -> None:
    """Apply column widths (inches) and font size to an existing table."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths):
                cell.width = Inches(widths[i])
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)


def _set_landscape_layout(doc: Document) -> None:
    """Rotate first section to landscape with tight margins for wide tables."""
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)


def write_doc(data: TableData) -> None:
    """
    Render :class:`TableData` and derived surrogate/stat tables into ``results.docx``.

    Czech headings and methodology paragraphs match the thesis write-up; source
    paths are listed in a footer paragraph for reproducibility.
    """
    cfg = load_config()
    from config_loader import dch_test_point_count

    mode_label = (
        f"TEST {dch_test_point_count()}" if dch_test_mode_from_env() else "FULL"
    )
    doc = Document()
    _set_landscape_layout(doc)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(8)
    doc.add_heading(f"Tabulka invariantů ({mode_label} běhy z /results)", level=1)
    # --- Table 1: per-coin invariant matrix ---
    doc.add_paragraph(
        "Zpoždění τ (řádek τ) je prodleva vložení z mutual information (``mutual.py`` → ``TAU_D2_<sym>``). "
        "τ_w je dekorelační čas z ``tau_w.py``. Theilerovo okno W (řádek W) se v tomto projektu "
        "**nastavuje rovno τ** (``W_D2_<sym> := TAU_D2_<sym>``); ACF/STP z ``theilers_w.bat`` jsou "
        "pouze diagnostické grafy."
    )

    rows = [
        row_values("Sample start", data.sample_start),
        row_values("Sample end", data.sample_end),
        row_values("N", data.n_vals),
        row_values("τ", data.tau),
        row_values("τw", data.tau_w),
        row_values("W", data.theiler_w),
        row_values("m", data.m_vals),
        row_values("d_c", data.dc),
        row_values("D_T (Takens)", data.takens),
        row_values("D_E (Ellner)", data.ellner),
        row_values("λmax", data.lle),
        row_values("RR", data.rr),
        row_values("DET", data.det),
        row_values("LAM", data.lam),
        row_values("MAXLINE", data.maxline),
        row_values("ENTR", data.entr),
        row_values("TT", data.tt),
        row_values("TREND", data.trend),
    ]

    table = doc.add_table(rows=1, cols=len(HEADERS))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(HEADERS):
        _set_cell_text(hdr[i], h, font_size=8.0, bold=True)

    for rv in rows:
        c = table.add_row().cells
        for i, v in enumerate(rv):
            _set_cell_text(c[i], v, font_size=8.0)

    _format_table(table, [1.25, 1.18, 1.18, 1.18, 1.18, 1.18, 1.18, 1.18], font_size=8.0)

    # --- Table 2: surrogate / bootstrap TS decisions ---
    doc.add_paragraph("")
    doc.add_heading(f"Výsledky surrogate testů ({mode_label})", level=1)
    doc.add_paragraph(
        "Stacionární bootstrap (B=100) původní řady poskytuje robustní střední hodnotu (boot_mean) "
        "a směrodatnou odchylku (boot_sd) invariantu. Pro jednu plně přemíchanou (reshuffle) řadu se "
        "spočítá tentýž invariant a vyhodnotí se testová statistika "
        "TS = (boot_mean − reshuffle) / boot_sd. Hypotéza H0 (řada je pouhý nezávislý šum) se zamítá, "
        "pokud |TS| > 3 — tedy existuje prokazatelná struktura/paměť coby předpoklad pro chaos, "
        "nikoli důkaz chaosu. Test je k dispozici pro TAKENS, ELLNER, LLE a (při zapnutém "
        "RQA bootstrapu v ``RQA.bat``) pro skalární RQA metriky; prázdné nebo ``nan`` buňky "
        "znamenají ``insufficient data`` nebo vypnutý bootstrap. "
        "Parametry embedování (τ, m) a Theilerovo okno W vstupují do ``hypothesis.py`` jako "
        "``--delay`` a ``--theiler`` (stejné W jako v TISEAN dávkových skriptech a PyRQA)."
    )

    sur_headers = [
        "Aktivum", "Invariant", "orig", "boot_mean", "boot_sd",
        "reshuffle", "TS", "|TS|", "Rozhodnutí",
    ]
    sur_table = doc.add_table(rows=1, cols=len(sur_headers))
    sur_table.style = "Table Grid"
    for i, h in enumerate(sur_headers):
        _set_cell_text(sur_table.rows[0].cells[i], h, font_size=7.5, bold=True)

    for row in _collect_surrogate_rows(cfg):
        cells = sur_table.add_row().cells
        vals = [
            row.symbol,
            row.invariant,
            row.orig,
            row.boot_mean,
            row.boot_sd,
            row.resh,
            row.ts,
            row.abs_ts,
            row.decision,
        ]
        for i, value in enumerate(vals):
            _set_cell_text(cells[i], value, font_size=7.0)

    _format_table(sur_table, [0.75, 0.95, 0.82, 0.82, 0.82, 0.82, 0.82, 0.82, 1.25], font_size=7.0)

    # --- Table 3: input series μ, σ (not invariant SDs) ---
    doc.add_paragraph("")
    doc.add_heading(f"Statistiky vstupních log-výnosových řad ({mode_label})", level=1)
    doc.add_paragraph(
        "Tato tabulka neobsahuje směrodatné odchylky invariantů. Uvádí průměr a výběrovou směrodatnou "
        "odchylku vstupních log-výnosových řad, ze kterých se invarianty počítají. Surogát je náhodná "
        "permutace originální řady; normal a t3.5 jsou referenční řady."
    )

    stat_headers = [
        "Aktivum",
        "orig μ",
        "orig σ",
        "surr μ",
        "surr σ",
        "normal μ",
        "normal σ",
        "t3.5 μ",
        "t3.5 σ",
    ]
    stat_table = doc.add_table(rows=1, cols=len(stat_headers))
    stat_table.style = "Table Grid"
    for i, h in enumerate(stat_headers):
        _set_cell_text(stat_table.rows[0].cells[i], h, font_size=6.8, bold=True)

    for row in _collect_series_stats_rows(cfg):
        cells = stat_table.add_row().cells
        vals = [
            row.symbol,
            row.orig_mean,
            row.orig_sd,
            row.surr_mean,
            row.surr_sd,
            row.normal_mean,
            row.normal_sd,
            row.t_mean,
            row.t_sd,
        ]
        for i, value in enumerate(vals):
            _set_cell_text(cells[i], value, font_size=6.5)

    _format_table(
        stat_table,
        [0.75, 0.9, 0.85, 0.9, 0.85, 0.95, 0.9, 0.9, 0.85],
        font_size=6.5,
    )

    # Footer: reproducibility paths (FULL vs test dirs from config_loader).
    p_dim = hypothesis_correlation_dimension_dir(cfg)
    p_lle = hypothesis_lambda_max_dir(cfg)
    p_rqa = hypothesis_rqa_dir(cfg)
    p_theiler = Path(theiler_summary_path(cfg))
    p_bat = Path(default_per_coin_settings_bat_path())
    doc.add_paragraph(
        f"Zdroj agregace hypothesis: {p_dim}, {p_lle}, {p_rqa}; "
        f"Theiler (souhrn z theilers_w.bat): {p_theiler} "
        f"(aktivní W_D2_* v {p_bat}); "
        f"tau_w, 2dc, mutual pod {Path(get_results_dir(cfg))}; log-výnosy z {get_data_dir(cfg)}."
    )
    out_path = _out_docx_path()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main() -> None:
    """CLI entry: ``collect`` → ``write_doc`` → print saved path."""
    cfg = load_config()
    write_doc(collect(cfg))
    out = _out_docx_path()
    print(f"Saved: {out}")


if __name__ == "__main__":
    main()

